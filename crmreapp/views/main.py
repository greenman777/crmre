#!/usr/bin/env python
# -*- coding: utf-8 -*-

from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from django.template import RequestContext
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.models import Group
from crmreauth.models import User
from django.db.models import Q
import json,os,unicodecsv
from django.core.files import File
from datetime import datetime
from django.template.loader import get_template

from crmre.utils import JSONResponse
from crmreapp import models
from django.http import HttpResponse

from docxtpl import DocxTemplate
from xhtml2pdf import pisa
from django.conf import settings
from django.template import Context, Template
from decimal import Decimal

from django.core.mail import EmailMessage

import cairosvg
from docx import Document
from docx.shared import Inches
from cStringIO import StringIO

from crmreapp import smssend
import xml.etree.ElementTree as etree
import re

from rapidsms.router import send, lookup_connections

from crmreapp.utils.num2t4ru import num2text, decimal2text
import urllib2

from crmreapp import serializers

import sys
reload(sys)
sys.setdefaultencoding('utf-8')

def generate_pdf(data,document):

    font_tag = u"""
    <!DOCTYPE HTML PUBLIC "-//W3C//DTD HTML 4.01 Transitional//EN" "http://www.w3.org/TR/html4/loose.dtd">
    <html>
    <head>
    <title>{{ article.title }}</title>
    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8" />
    <style type="text/css">
        @page { size: A4; margin: 1cm; }
        @font-face { font-family: Arial; src: url("{{MEDIA_ROOT}}/font/arial.ttf"); }
        @font-face { font-family: Arial; font-weight: bold; src: url("{{MEDIA_ROOT}}/font/arialbd.ttf"); }
        @font-face { font-family: Arial; font-style: italic;src: url("{{MEDIA_ROOT}}/font/ariali.ttf"); }
        @font-face { font-family: Arial; font-weight: bold; font-style: italic;src: url("{{MEDIA_ROOT}}/font/arialbi.ttf"); }
        @font-face { font-family: Times New Roman; src: url("{{MEDIA_ROOT}}/font/times.ttf"); }
        @font-face { font-family: Times New Roman; font-weight: bold; src: url("{{MEDIA_ROOT}}/font/timesbd.ttf"); }
        @font-face { font-family: Times New Roman; font-style: italic;src: url("{{MEDIA_ROOT}}/font/timesi.ttf"); }
        @font-face { font-family: Times New Roman; font-weight: bold; font-style: italic;src: url("{{MEDIA_ROOT}}/font/timesbi.ttf"); }
        body { font-family: Arial; }
    </style>
    </head>
    <body>
    """
    template = Template(font_tag + document.content+"</body></html>")
    html = template.render(Context(data))
    result = StringIO()
    pdf = pisa.pisaDocument(StringIO(html.encode('utf-8')), result, encoding='UTF-8')
    if not pdf.err:
        return (result.getvalue(), 'application/pdf')
    return False

@csrf_exempt
def upload_avito(request):
    f = open(os.path.join(settings.MEDIA_ROOT,"uploading","uploading_avito_rn43.xml"), 'r')
    xml = File(f)
    response = HttpResponse(xml,content_type="text/xml")
    response['Content-Disposidescrtion']='attachment; filename="uploading_avito_rn43.xml"'
    return response

@csrf_exempt
def upload_yandex(request):
    f = open(os.path.join(settings.MEDIA_ROOT,"uploading","uploading_yandex_rn43.xml"), 'r')
    xml = File(f)
    response = HttpResponse(xml,content_type="text/xml")
    response['Content-Disposition']='attachment; filename="uploading_yandex_rn43.xml"'
    return response

@csrf_exempt
def upload_portal(request):
    f = open(os.path.join(settings.MEDIA_ROOT,"uploading","uploading_portal_rn43.xml"), 'r')
    xml = File(f)
    response = HttpResponse(xml,content_type="text/xml")
    response['Content-Disposition']='attachment; filename="uploading_portal_rn43.xml"'
    return response

@csrf_exempt
def request_portal(request):
    #try:
    if request.method == 'POST':
        xml = request.body
        root = etree.fromstring(xml)
        notice_type = root.find('NoticeType').text
        datetime = root.find('DateTime').text
        user_name = root.find('UserName').text
        user_phone = root.find('UserPhone').text
        managers = []
        if (notice_type=='object'):
            type = root.find('Type').text
            object_id_crm = root.find('ObjectIdCrm').text
            manager_id_crm = root.find('ManagerIdCrm').text
            managers = [User.objects.get(pk=int(manager_id_crm))]
            content = unicode(" ".join((type, u", объект", object_id_crm, u", клиент", user_name, user_phone)))
        elif (notice_type=='other'):
            title = root.find('Title').text
            managers = User.objects.filter(user_permissions__codename='callback_manager').distinct()
            content = unicode(" ".join((title, u", клиент", user_name, user_phone)))
        success_send = False
        for manager in managers:
            if (manager.sms_notification and (len(manager.phone)==11)):
                phone = "".join(("7",manager.phone[1:]))
                connections = lookup_connections(backend="kannel-beeline-smpp",identities=[phone])
                send(content, connections=connections)
                success_send = True
        if success_send:
            response = JSONResponse({'success':True,'messages':content})
        else:
            response = JSONResponse({'success':False,'messages':'Not send a message to'})
    else:
        response = JSONResponse({'success':False,'messages':'Method is not supported'})
    return response

@csrf_exempt
def test_portal(request):
    xml_data = """
                <?xml version="1.0" encoding="utf-8" ?>
                <Request>
                    <Type>Купить с ипотекой</Type>
                    <NoticeType>object</NoticeType>
                    <DateTime>2015-04-29 12:14:06</DateTime>
                    <ObjectIdCrm>349154735</ObjectIdCrm>
                    <ManagerIdCrm>1</ManagerIdCrm>
                    <UserName>Дмитрий</UserName>
                    <UserPhone>89127113553</UserPhone>
                  </Request>
                """
    url = "http://127.0.0.1:8080/request_portal/"
    req = urllib2.Request(url=url, data=xml_data)
    f = urllib2.urlopen(req)
    #r = requests.post(url,data=xml_data, headers={'Content-Type': 'text/xml'},allow_redirects=False)
    response = JSONResponse({'success':True,'messages':'Ok'})
    return response

@csrf_exempt
def send_messages(request):

    if request.method == 'POST':
        getparams = request.POST.copy()
        messages_arr = json.loads(request.POST['messages_arr'])
        messages_type = getparams.get('messages_type')
        pkg_msg = {models.MessageType.objects.get(name=u'Подписка 1').id:Q(pkg_sms_1=True),
                   models.MessageType.objects.get(name=u'Подписка 2').id:Q(pkg_sms_2=True),
                   models.MessageType.objects.get(name=u'Подписка 3').id:Q(pkg_sms_3=True),
                   models.MessageType.objects.get(name=u'Подписка 4').id:Q(pkg_sms_4=True),
                   models.MessageType.objects.get(name=u'Подписка 5').id:Q(pkg_sms_5=True)}
        content = ""
        if messages_type == 'sms':
            sms = smssend.AtomParkSMS()
            for messages in messages_arr:
                phones = ["".join(("7",phone[0][1:])) for phone in models.Clients.objects.filter(pkg_msg[messages[0]]).values_list('phone_represent') if len(phone[0]) == 11]
                content = "<br>".join((content,u"Сообщение " + unicode(messages[0]) + ": " + sms.send_sms(messages[1], phones)))
            response = JSONResponse({'success':True,'messages':content})
        if messages_type == 'email':
            for messages in messages_arr:
                email = [str(email[0]) for email in models.Clients.objects.filter(pkg_msg[messages[0]]).values_list('email')]
                if (email):
                    try:
                        msg = EmailMessage(u'Новое сообщение для Вас',messages[1],'noreplay@rn43.ru',email)
                        msg.content_subtype = "html"
                        msg.send()
                        content_messsage = u"Сообщения успешно отправлены: " + unicode(len(email)) + u" получателей"
                    except:
                        content_messsage = u"При отправке сообщений произошли ошибки"
                else:
                    content_messsage = u"Нет ни одного подписчика на сообщение"
                content = "<br>".join((content,u"Сообщение " + unicode(messages[0]) + ": " + content_messsage))
            response = JSONResponse({'success':True,'messages':content})
        return response

def generate_docx(data):

    template = get_template('docx.html')
    result = StringIO()
    document = Document()
    document.add_heading(data['type_uploading'], 0)
    objects = data['object_list']
    content = {}
    for obj in objects:
        content['obj']=obj
        render = template.render(Context(content))
        document.add_paragraph(render+"\n"+100*'_')
        for photo in obj.photos_set.all().order_by('description')[0:1]:
            try:
                document.add_picture(photo.photo, width=Inches(1.25))
            except:
                pass
        #document.add_paragraph(100*'_',style='Normal')
    document.save(result)
    length = result.tell()
    result.seek(0)
    return (result.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',length)

@csrf_exempt
def send_offer(request):
    if request.method == 'POST':
        model = models.DocumentTemplates
        arr_offer = json.loads(request.POST['arr_offer'])
        offer_dict = {}
        failure_send_clients = []
        for offer in arr_offer:
            if offer_dict.get(offer[0]):
                offer_dict[offer[0]].append(offer[1])
            else:
                offer_dict[offer[0]] = []
                offer_dict[offer[0]].append(offer[1])
        for key in offer_dict:
            data = {}
            order_buy = models.OrdersBuy.objects.get(id=key)
            order_sale_list = [models.OrdersSale.objects.get(id=value) for value in offer_dict[key] ]
            data['client'] = order_buy.client
            data['user'] = order_buy.performer
            data['object_list'] = order_sale_list
            data['MEDIA_ROOT'] = settings.MEDIA_ROOT
            data['date'] = datetime.now()
            document = model.objects.get(name=u"Предложения для клиента")
            (result,content_type) = generate_pdf(data,document)
            to_email = data['client'].email
            if (to_email):
                try:
                    message = EmailMessage(u'Новые предложения по вашей заявке',
                                           u"""Новые предложения по вашей заявке представлены во вложении к письму.
                                            Письмо подготовлено и отправлено автоматически. Пожалуйста, не отвечайте на это письмо.
                                            Для обратной связи с нами воспользуйтесь реквизитами, указанными в письме""",
                                           'noreplay@rn43.ru',[to_email])
                    message.attach(u"Новые предложения.pdf", result, content_type)
                    message.send();
                except:
                    failure_send_clients.append(data['client'].client_name)
            else:
                failure_send_clients.append(data['client'].client_name)
        response = JSONResponse({'success':True,'messages':failure_send_clients})
        return response

@csrf_exempt
def open_offer(request):
    if request.method == 'POST':
        model = models.DocumentTemplates
        arr_offer = json.loads(request.POST['arr_offer'])
        offer_dict = {}
        for offer in arr_offer:
            if offer_dict.get(offer[0]):
                offer_dict[offer[0]].append(offer[1])
            else:
                offer_dict[offer[0]] = []
                offer_dict[offer[0]].append(offer[1])
        for key in offer_dict:
            data = {}
            order_buy = models.OrdersBuy.objects.get(id=key)
            order_sale_list = [models.OrdersSale.objects.get(id=value) for value in offer_dict[key] ]
            data['client_type'] = order_buy.client.client_type
            if ((order_buy.performer.id == request.user.id) or request.user.has_perm('crmreapp.view_hidden_fields_clients')):
                data['client'] = order_buy.client
            else:
                data['client'] = ""
            data['user'] = order_buy.performer
            data['object_list'] = order_sale_list
            data['MEDIA_ROOT'] = settings.MEDIA_ROOT
            data['date'] = datetime.now()
            document = model.objects.get(name=u"Предложения для клиента")
            (result,content_type) = generate_pdf(data,document)
            response = HttpResponse(result, content_type=content_type)
            response['Content-Disposition'] = 'attachment'
            return response


def change_field(value, item):
    if item == 'planishing__name' and value:
        value = " ".join((unicode(value),u"планировка"))
    if item == 'condition__name' and value:
        value = " ".join((u"состояние",unicode(value)))
    if item == 'performer__phone' and len(value) > 0:
        value = "8"+value[1:]
    if value is None:
        value = ""
    if item in ('floors', 'floor'):
        value = "'"+unicode(value)+"'"
    if type(value) is Decimal:
        value = unicode(float(value)).replace('.', ',')
    return unicode(value)


@csrf_exempt
def export_csv(data,data_headers):
    response = HttpResponse(content_type='text/csv')
    response['success'] = True
    response['Content-Disposition'] = 'attachment; filename=export_data.csv'
    writer = unicodecsv.writer(response,delimiter=';', encoding='cp1251')
    if len(data) > 0:
        headers = [header['name'] for header in data_headers]
        writer.writerow(headers)
        for row in data:
            fields = []
            for header in data_headers:
                if header.get('items') == ['description']:
                    p = re.compile(r'<.*?>')
                    row['description'] = p.sub('', row['description'])
                fields.append(header['separator'].join((change_field(row[item], item) for item in header['items'])))
            try:
                writer.writerow(fields)
            except UnicodeEncodeError:
                response['success'] = False
    return response

@csrf_exempt
def documents_upload(request):
    if request.method == 'POST':
        name = request.POST['name']
        offer = models.Offer.objects.get(pk=request.POST['offer'])
        newfile = models.Documents(offer = offer, name = name, file = request.FILES['file'])
        newfile.save()
        response = JSONResponse({'success':True,'message': u'Файл загружен!'})
        return response

@csrf_exempt
def uploading_templatesdoc(request):
    if request.method == 'POST':
        getparams = request.POST.copy()
        order_type = getparams.get('order_type')
        order_id = getparams.get('order')
        template = getparams.get('template')
        male_units = ((u'рубль', u'рубля', u'рублей'), 'm')
        operation_complet = models.ResultOperation.objects.get(name=u"успешная")
        operation_registr = models.OperationType.objects.get(name=u"4. регистрация")
        doc = DocxTemplate(models.TemplatesDoc.objects.get(pk=template).file)
        if (order_type == 'sale'):
            order = models.OrdersSale.objects.get(pk=order_id)
            try:
                price = int(order.price)
                number_rooms = order.number_rooms
                address_object = " ".join((order.city.name,order.street.name,order.house_number))
                floor = order.floor
                total_space = order.total_space
                living_space = order.living_space
                object_type = order.object_type.name
            except Exception as err:
                price = 0
                number_rooms = 0
                address_object = ""
                floor = 0
                total_space = 0
                living_space = 0
                object_type = ""
        else:
            order = models.OrdersBuy.objects.get(pk=order_id)
            try:
                price = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.price)
                number_rooms = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.number_rooms)
                city = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.city.name)
                street = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.street.name)
                house_number = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.house_number)
                address_object = " ".join((city,street,house_number))
                floor = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.floor)
                total_space = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.total_space)
                living_space = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.living_space)
                object_type = int(order.offer_set.filter(hystoryservice__operation=operation_registr,hystoryservice__result_operation=operation_complet)[0].order_sale.object_type.name)
            except Exception as err:
                price = 0
                number_rooms = 0
                address_object = ""
                floor = 0
                total_space = 0
                living_space = 0
                object_type = ""
        context = { 'client_name':unicode(order.client.client_name),  #имя клиента
                    'price': price,  #цена объекта
                    'price_text': unicode(num2text(price, male_units)),  #цена объекта текстом
                    'contract_number': order.contract_number,  #номер договора
                    'date_birth': order.client.date_birth.strftime('%d.%m.%Y') if order.client.date_birth else "", #дата рождения
                    'passport_series':order.client.passport_series,  #серия паспорта
                    'passport_number': order.client.passport_number,  # номер паспорта
                    'date_issue': order.client.date_issue.strftime('%d.%m.%Y') if order.client.date_issue else "", # дата выдачи паспорта
                    'output_place': order.client.output_place,  # место выдачи паспорта
                    'address_registr': order.client.address_registr,  # адрес регистрации клиента
                    'address_actual': order.client.address_actual,  # фактический адрес клиента
                    'email': order.client.email,  # email клиента
                    'number_rooms': number_rooms,  # количество комнат
                    'address_object': address_object,  # адрес объекта
                    'floor': floor,  # этаж
                    'phone_main': order.client.phone_main,  # телефон клиента
                    'total_space': total_space,  # общая площадь
                    'living_space': living_space,  # жилая площадь
                    'object_type': object_type,  # тип объекта
                    }
        doc.render(context)
        result = StringIO()
        doc.save(result)
        length = result.tell()
        result.seek(0)
        result.getvalue()
        response = HttpResponse(result,content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition']='attachment; filename="uploading.docx"'
        response['Content-Length'] = length
        return response

@csrf_exempt
def uploading(request):
    if request.method == 'POST':
        getparams = request.POST.copy()
        status_activ = models.OrderStatus.objects.get(name=u"активная")
        objects_new = models.ObjectType.objects.filter(name=u"квартира в новостройке")
        type_uploading =  getparams.get('type_uploading')
        data = {}
        data['MEDIA_ROOT'] = settings.MEDIA_ROOT
        if (type_uploading == u"Выгрузка в газету ЯЖ"):
            data['type_uploading'] = type_uploading
            data['object_list'] = models.OrdersSale.objects.filter(classified_resources=True,status=status_activ).exclude(object_type__in=objects_new).order_by('object_category__name','object_type__name','number_rooms','city__name','microdistrict__name','street__name')
            (result,content_type,length) = generate_docx(data)
            response = HttpResponse(result,content_type=content_type)
            response['Content-Disposition']='attachment; filename="uploading.docx"'
            response['Content-Length'] = length
            return response
        elif (type_uploading == u"Выгрузка в АН"):
            data_headers = [{'name':u'город н/п','separator':'','items':['city__name']},
                            {'name':u'улица, номер дома','separator':', ','items':['street__name','house_number']},
                            {'name':u'район','separator':'','items':['district__name']},
                            {'name':u'мкр район','separator':'','items':['microdistrict__name']},
                            {'name':u'тип недвижимости','separator':'','items':['object_type__name']},
                            {'name':u'количество комнат','separator':'','items':['number_rooms']},
                            {'name':u'площадь','separator':'/','items':['total_space','living_space','kitchen_space']},
                            {'name':u'этажей/этаж','separator':'/','items':['floors','floor']},
                            {'name':u'планировка','separator':', ','items':['material_walls__name','planishing__name','balcony__name','condition__name']},
                            {'name':u'цена','separator':'','items':['price']},
                            {'name':u'описание','separator':'','items':['description']},
                            {'name':u'агент','separator':' ','items':['performer__last_name','performer__first_name']},
                            {'name':u'телефон','separator':'','items':['performer__phone']},]
            contract_exl = models.ContractType.objects.get(name=u"эксклюзивный")
            filter_Q = Q(other_agency = True)|Q(contract_type=contract_exl,contract_number__isnull=False,contract_date__isnull=False)
            fields = []
            for field in (header['items'] for header in data_headers):
                fields.extend(field)
            data = models.OrdersSale.objects.filter(filter_Q).filter(status=status_activ).exclude(object_type__in=objects_new).order_by('object_category__name','object_type__name','number_rooms','city__name','microdistrict__name','street__name').values(*fields)
            return export_csv(data,data_headers)

@csrf_exempt
def uploading_orders(request):
    if request.method == 'POST':
        getparams = request.POST.copy()
        date_start = datetime.strptime(getparams.get('date_start'), "%Y-%m-%d")
        date_stop = datetime.strptime(getparams.get('date_stop'), "%Y-%m-%d")
        order_status = getparams.getlist('order_status')
        order_type = int(str(getparams.get('order_type')))
        order_status = [int(str(status)) for status in order_status]
        data = {}
        data['MEDIA_ROOT'] = settings.MEDIA_ROOT
        if order_type == 1:
            data_headers = [
            {'name': u'номер заявки', 'separator': '', 'items': ['index']},
            {'name': u'заголовок', 'separator': '', 'items': ['heading']},
            {'name': u'тип объекта', 'separator': '', 'items': ['object_type__name']},
            {'name': u'дата создания', 'separator': '', 'items': ['create_date']},
            {'name':u'город н/п','separator':'','items':['city__name']},
            {'name':u'улица','separator':', ','items':['street__name']},
            {'name': u'номер дома', 'separator': ', ', 'items': ['house_number']},
            {'name': u'квартира', 'separator': ', ', 'items': ['house_apartment']},
            {'name':u'цена','separator':'','items':['price']},
            {'name': u'статус','separator': '','items':['status__name']},
            {'name':u'агент','separator':' ','items':['performer__last_name','performer__first_name']},
            {'name':u'номер клиента','separator':'','items':['client__index']},
            {'name': u'представитель', 'separator': '', 'items': ['client__represent']},
            {'name': u'телефон', 'separator': '', 'items': ['client__phone_represent']},
            {'name': u'клиент', 'separator': '', 'items': ['client__client_name']},
            {'name': u'тип клиента', 'separator': '', 'items': ['client__client_type']},
            {'name': u'email', 'separator': '', 'items': ['client__email']},
            {'name': u'рекл.бесплатная', 'separator': '', 'items': ['classified_resources']},
            {'name': u'рекл.платная', 'separator': '', 'items': ['toll_resources']},
            {'name': u'рекл.горячее', 'separator': '', 'items': ['hot_offer']},
            {'name': u'рекл.элитное', 'separator': '', 'items': ['luxury_housing']},]
        elif order_type == 2:
            data_headers = [
            {'name': u'номер заявки', 'separator': '', 'items': ['index']},
            {'name': u'заголовок', 'separator': '', 'items': ['heading']},
            {'name': u'дата создания', 'separator': '', 'items': ['create_date']},
            {'name': u'статус', 'separator': '', 'items': ['status__name']},
            {'name': u'агент', 'separator': ' ', 'items': ['performer__last_name', 'performer__first_name']},
            {'name': u'номер клиента', 'separator': '', 'items': ['client__index']},
            {'name': u'представитель', 'separator': '', 'items': ['client__represent']},
            {'name': u'телефон', 'separator': '', 'items': ['client__phone_represent']},
            {'name': u'клиент', 'separator': '', 'items': ['client__client_name']},
            {'name': u'тип клиента', 'separator': '', 'items': ['client__client_type']},
            {'name': u'email', 'separator': '', 'items': ['client__email']},
            {'name': u'рекл.бесплатная', 'separator': '', 'items': ['classified_resources']},
            {'name': u'рекл.платная', 'separator': '', 'items': ['toll_resources']},
            ]
        else:
            data_headers = []
        fields = []
        for field in (header['items'] for header in data_headers):
            fields.extend(field)
        if order_type==1:
            data = models.OrdersSale.objects.filter(status__in=order_status,create_date__range=(date_start,date_stop)).order_by('create_date', 'object_type__name', 'number_rooms','city__name', 'street__name').values(*fields)
        elif order_type==2:
            data = models.OrdersBuy.objects.filter(status__in=order_status,create_date__range=(date_start,date_stop)).order_by('create_date').values(*fields)
        else:
            data = []
        return export_csv(data,data_headers)

@login_required
def mainpage(request):

    user_groups = list(request.user.groups.values_list('id',flat=True))
    return render(request,'index.html',{'user_groups':user_groups})

def photo_upload(request):
    if request.method == 'POST':
        photo_index = 0
        for description in request.POST['description'].split(','):
            obj = models.OrdersSale.objects.get(pk=request.POST['object'])
            newphoto = models.Photos(object = obj, description = description.split('.')[0], photo = request.FILES.getlist('photo')[photo_index])
            newphoto.save()
            photo_index = photo_index + 1;
        response = JSONResponse({'success':True,'message': u'Фотографии загружены!'})
        return response

def plan_photo_upload(request):
    if request.method == 'POST':
        description = request.POST['description']
        plan = models.Plan.objects.get(pk=request.POST['plan'])
        newphoto = models.PlanPhotos(plan = plan, description = description, photo = request.FILES['photo'])
        newphoto.save()
        response = JSONResponse({'success':True,'message': u'Фотография загружена!'})
        return response

def building_photo_upload(request):
    if request.method == 'POST':
        description = request.POST['description']
        building = models.Buildings.objects.get(pk=request.POST['building'])
        newphoto = models.BuildingPhotos(building = building, description = description, photo = request.FILES['photo'])
        newphoto.save()
        response = JSONResponse({'success':True,'message': u'Фотография загружена!'})
        return response

"""
Поиск объектов для заявки (спрос)
"""
def offer_news(request):
    if request.method == 'GET':
        getparams = request.GET.copy()
        order_id =  getparams.get('order_id')
        order_type =  getparams.get('order_type')
        search_my = json.loads(getparams.get('search_my'))
        if order_id != None:
            filter_Q = Q()
            user_groups = request.user.groups.values_list('name',flat=True)
            if (u"_Внешние агенты" in user_groups) :
                filter_Q = filter_Q & Q(other_agency = True)
            if order_type == 'appOfferBuyList':
                """
                Поиск объектов для заявки (спрос)
                """
                #заявка на спрос
                object_search = models.OrdersBuy.objects.get(pk=order_id)#объект по которому ищем

                if search_my:
                    filter_Q = filter_Q & Q(performer = object_search.performer)

                offer_my = models.Offer.objects.filter(order_buy=order_id)#список предложений
                offer_new = [dict(element) for element in serializers.OfferSerializer(offer_my, many=True).data]

                object_type_list = [obj.object_type for obj in object_search.listobjecttype_set.all()]#список типов объектов для поиска

                if len(object_type_list):
                    filter_Q = filter_Q & (Q(object_type__in = object_type_list)|Q(object_type__isnull=True))

                district_list = [obj.district for obj in object_search.listdistrict_set.all()]#список районов для поиска
                if len(district_list):
                    filter_Q = filter_Q & (Q(district__in = district_list)|Q(district__isnull=True))

                city_list = [obj.city for obj in object_search.listcity_set.all()]#список городов для поиска
                if len(city_list):
                    filter_Q = filter_Q & (Q(city__in = city_list)|Q(city__isnull=True))

                microdistrict_list = [obj.microdistrict for obj in object_search.listmicrodistrict_set.all()]#список микрорайонов для поиска
                if len(microdistrict_list):
                    filter_Q = filter_Q & (Q(microdistrict__in = microdistrict_list)|Q(microdistrict__isnull=True))

                street_list = [obj.street for obj in object_search.liststreet_set.all()]#список улиц для поиска
                if len(street_list):
                    filter_Q = filter_Q & (Q(street__in = street_list)|Q(street__isnull=True))

                rooms_list = [obj.number_rooms for obj in object_search.listrooms_set.all()]#список комнат для поиска
                if len(rooms_list):
                    filter_Q = filter_Q & (Q(number_rooms__in = rooms_list)|Q(number_rooms__isnull=True)|Q(number_rooms=0))

                price_from = object_search.price_from
                price_to = object_search.price_to
                if (price_from and price_to):
                    filter_Q = filter_Q & (Q(price__range = (price_from,price_to))|Q(price__isnull=True)|Q(price=0))
                elif (price_from):
                    filter_Q = filter_Q & (Q(price__gte = price_from)|Q(price__isnull=True)|Q(price=0))
                elif (price_to):
                    filter_Q = filter_Q & (Q(price__lte = price_to)|Q(price__isnull=True)|Q(price=0))

                space_from = object_search.space_from
                space_to = object_search.space_to
                if (space_from and space_to):
                    filter_Q = filter_Q & (Q(total_space__range = (space_from,space_to))|Q(total_space__isnull=True)|Q(total_space=0))
                elif (space_from):
                    filter_Q = filter_Q & (Q(total_space__gte = space_from)|Q(total_space__isnull=True)|Q(total_space=0))
                elif (space_to):
                    filter_Q = filter_Q & (Q(total_space__lte = space_to)|Q(total_space__isnull=True)|Q(total_space=0))

                remoteness_from = object_search.remoteness_from
                remoteness_to = object_search.remoteness_to
                if (remoteness_from and remoteness_to):
                    filter_Q = filter_Q & (Q(remoteness_center__range = (remoteness_from,remoteness_to))|Q(remoteness_center__isnull=True))
                elif (remoteness_from):
                    filter_Q = filter_Q & (Q(remoteness_center__gte = remoteness_from)|Q(remoteness_center__isnull=True))
                elif (remoteness_to):
                    filter_Q = filter_Q & (Q(remoteness_center__lte = remoteness_to)|Q(remoteness_center__isnull=True))

                #ищем заявки на предложение с подходящими условиями
                orders_sale = models.OrdersSale.objects.filter(filter_Q,
                                                               status = models.OrderStatus.objects.get(name=u'активная'),
                                                               transaction_type = object_search.transaction_type,
                                                               object_category = object_search.object_category
                                                               )
                #перебираем все найденные объекты и если их нет в предложении, добавляем
                for order in orders_sale:
                    if offer_my.filter(order_sale=order.id).count() <= 0:
                        offer_new += [{
                            "order_buy": object_search.id, "order_sale": order.id,
                            "informed": False, "stage": 1,
                            "order_sale_client_name": order.client.client_name,
                            "order_sale_heading":order.heading,
                            "order_sale_street_name": order.street.name if order.street else "",
                            "order_sale_house_number": order.house_number,
                            "order_sale_microdistrict_name": order.microdistrict.name if order.microdistrict else "",
                            "order_sale_city_name": order.city.name if order.city else "",
                            "order_sale_space": order.total_space,
                            "order_sale_price": order.price,
                            "order_sale_object_type_name": order.object_type.name if order.object_type else "",
                            "order_buy_performer_id": order.performer.id,
                            "order_sale_performer_name": " ".join((order.performer.last_name, order.performer.first_name)),
                            "order_sale_create_date": order.create_date,
                            "order_sale_index": order.index,
                        }]

            if order_type == 'appOfferSaleList':
                """
                Поиск объектов для заявки (предложение)
                """
                #заявка на спрос
                object_search = models.OrdersSale.objects.get(pk=order_id)#объект по которому ищем

                if search_my:
                    filter_Q = filter_Q & Q(performer = object_search.performer)

                offer_my = models.Offer.objects.filter(order_sale=order_id)#список предложений
                offer_new = [dict(element) for element in serializers.OfferSerializer(offer_my, many=True).data]

                price = object_search.price
                if (price):
                    filter_Q = filter_Q & (Q(price_from__lte = price)|Q(price_from__isnull=True)|Q(price_from=0)) & (Q(price_to__gte = price)|Q(price_to__isnull=True)|Q(price_to=0))

                space = object_search.total_space
                if (space):
                    filter_Q = filter_Q & (Q(space_from__lte = space)|Q(space_from__isnull=True)|Q(space_from=0)) & (Q(space_to__gte = space)|Q(space_to__isnull=True)|Q(space_to=0))

                remoteness = object_search.remoteness_center
                if (remoteness):
                    filter_Q = filter_Q & (Q(remoteness_from__lte = remoteness)|Q(remoteness_from__isnull=True)) & (Q(remoteness_to__gte = remoteness)|Q(remoteness_to__isnull=True))

                object_type = object_search.object_type
                if (object_type):
                    filter_Q = filter_Q & (Q(listobjecttype__object_type = object_type)|Q(listobjecttype__object_type__isnull=True))

                district = object_search.district
                if (district):
                    filter_Q = filter_Q & (Q(listdistrict__district = district)|Q(listdistrict__district__isnull=True))

                microdistrict = object_search.microdistrict
                if (microdistrict):
                    filter_Q = filter_Q & (Q(listmicrodistrict__microdistrict = microdistrict)|Q(listmicrodistrict__microdistrict__isnull=True))

                city = object_search.city
                if (city):
                    filter_Q = filter_Q & (Q(listcity__city = city)|Q(listcity__city__isnull=True))

                street = object_search.street
                if (street):
                    filter_Q = filter_Q & (Q(liststreet__street = street)|Q(liststreet__street__isnull=True))

                number_rooms = object_search.number_rooms
                if (street):
                    filter_Q = filter_Q & (Q(listrooms__number_rooms = number_rooms)|Q(listrooms__number_rooms__isnull=True))
                #ищем заявки на предложение с подходящими условиями
                orders_buy = models.OrdersBuy.objects.filter(
                    filter_Q,
                    status = models.OrderStatus.objects.get(name=u'активная'),
                    transaction_type = object_search.transaction_type,
                    object_category = object_search.object_category)
                #перебираем все найденные объекты и если их нет в предложении, добавляем
                for order in orders_buy:
                    if offer_my.filter(order_buy=order.id).count() <= 0:
                        offer_new += [{
                            "order_sale": object_search.id, "order_buy": order.id,
                            "informed": False, "stage": 1,
                            "order_buy_client_name": order.client.client_name,
                            "order_buy_heading":order.heading,
                            "order_buy_space": (" ".join(('от', str(order.space_from))) if order.space_from else "") +
                                               (" ".join(('до', str(order.space_to))) if order.space_to else ""),
                            "order_buy_price": (" ".join(('от', str(order.price_from))) if order.price_from else "") +
                                               (" ".join(('до', str(order.price_to))) if order.price_to else ""),
                            "order_buy_performer_id": order.performer.id,
                            "order_buy_performer_name": " ".join((order.performer.last_name, order.performer.first_name)),
                            "order_buy_create_date": order.create_date,
                            "order_buy_index": order.index,
                        }]

            return JSONResponse(offer_new)


def send_notification(request):
    if request.method == 'POST':
        getparams = request.POST.copy()
        recipients =  json.loads(getparams.get('recipients'))
        message =  getparams.get('message')
        sendsms = json.loads(getparams.get('sendsms'))
        if len(recipients) == 0:
            recipients = User.objects.filter(is_active=True)
        else:
            recipients = [User.objects.get(id=recipient) for recipient in recipients]
        for recipient in recipients:
            models.Notifications.objects.create(user=recipient,sender=request.user,message=message,sendsms=sendsms)
        response = JSONResponse({'success':True,'messages':u'Сообщение успешно отправлено!'})
        return response

def notifications_news(request):
    model = models.Notifications
    if request.method == 'POST':
        getparams = request.POST.copy()
        user =  getparams.get('user')
        if user != None:
            free_status = models.OrderStatus.objects.get(name=u'свободная')
            complet_status = models.OrderStatus.objects.get(name=u'сделка завершена')
            close_status = models.OrderStatus.objects.get(name=u'отказная')
            count_free_buy = models.OrdersBuy.objects.filter(status = free_status).count()
            count_free_sale = models.OrdersSale.objects.filter(status = free_status).count()
            count_complet_buy = models.OrdersBuy.objects.filter(status = complet_status).count()
            count_complet_sale = models.OrdersSale.objects.filter(status = complet_status).count()
            count_close_buy = models.OrdersBuy.objects.filter(status = close_status).count()
            count_close_sale = models.OrdersSale.objects.filter(status = close_status).count()
            count_news = model.objects.filter(user=user,read=False).exclude(sender=user).count()
            response = JSONResponse({'success':True,'messages':{'count_news':count_news,
                                                                'count_free_buy':count_free_buy,
                                                                'count_free_sale':count_free_sale,
                                                                'count_complet_buy':count_complet_buy+count_close_buy,
                                                                'count_complet_sale':count_complet_sale+count_close_sale,}})
            return response

def count_orders(request):
    model_sale = models.OrdersSale
    model_buy = models.OrdersBuy
    if request.method == 'POST':
        getparams = request.POST.copy()
        user =  getparams.get('user')
        if user != None:
            count_orders_sale = model_sale.objects.filter(performer=user,status=models.OrderStatus.objects.get(name=u'активная')).count()
            count_orders_buy = model_buy.objects.filter(performer=user,status=models.OrderStatus.objects.get(name=u'активная')).count()
            count_orders = count_orders_sale + count_orders_buy
            response = JSONResponse({'success':True,'messages':{'count_orders':count_orders}})
            return response

def credit_manager(request):
    model_user = User
    for user in model_user.objects.all():
        if user.has_perm('crmreapp.credit_manager') and (not user.is_superuser):
            response = JSONResponse({'success':True,'messages':{'credit_manager':user.id}})
            return response
    response = JSONResponse({'success':True,'messages':{'credit_manager':0}})
    return response

def manager(request):
    model_user = User
    if request.method == 'POST':
        getparams = request.POST.copy()
        agent =  getparams.get('agent')
        if agent != None:
            user = model_user.objects.get(id=int(agent))
            if not user.is_superuser:
                group_managers = Group.objects.get(name=u'_Руководители групп')
                managers_list = list(group_managers.user_set.all())
                group_boss = Group.objects.get(name=u'_Высший менеджмент')
                boss_list = list(group_boss.user_set.values_list('id'))
                recipients = [val[0] for val in boss_list]
                user_groups_set = set([val[0] for val in user.groups.values_list('id')])
                for manager in managers_list:
                    manager_groups_set = set([val[0] for val in manager.groups.values_list('id')])
                    if set(user_groups_set) & set(manager_groups_set):
                        recipients.append(manager.id)
                response = JSONResponse({'success':True,'messages':{'managers':recipients}})
                return response
    response = JSONResponse({'success':True,'messages':{'managers':0}})
    return response

@csrf_exempt
def svg_to_png(request):
    if request.method == 'POST':
        svg_code = request.POST['svg']
        result = StringIO()
        cairosvg.svg2png(bytestring=svg_code.encode('utf-8', 'ignore'),write_to=result)
        response = HttpResponse(result.getvalue(),content_type='image/png')
        response['Content-Disposition']='attachment; filename="chart.png"'
        return response
    else:
        response = JSONResponse({'success': False, 'messages': {'Error'}})
        return response